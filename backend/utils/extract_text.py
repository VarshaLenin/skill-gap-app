"""Extract plain text from an uploaded resume file (.pdf, .docx, .txt)."""
import io

import pdfplumber
from docx import Document


class UnsupportedFileType(Exception):
    pass


def extract_text(file_storage) -> str:
    """
    Extract text from a Flask FileStorage object based on its filename extension.
    """
    filename = (file_storage.filename or "").lower()
    raw_bytes = file_storage.read()

    if filename.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    if filename.endswith(".docx"):
        return _extract_docx(raw_bytes)
    if filename.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore").strip()

    raise UnsupportedFileType(
        "Unsupported file type. Please upload a .pdf, .docx, or .txt resume."
    )


def _extract_pdf(raw_bytes: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts).strip()


def _extract_docx(raw_bytes: bytes) -> str:
    document = Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]

    # Also pull text out of any tables (resumes sometimes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs).strip()
